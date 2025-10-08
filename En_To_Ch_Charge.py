"""
Word文档英译中翻译器
使用 python-docx 和 googletrans 实现英文Word文档翻译
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from docx import Document
from docx.oxml import OxmlElement
from googletrans import Translator
import threading
import os
from datetime import datetime
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
import queue


class WordTranslatorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Word文档英译中翻译器")
        self.root.geometry("900x700")
        self.root.resizable(True, True)

        # 变量初始化
        self.input_file = None
        self.translated_doc = None
        self.translator = Translator()
        self.is_translating = False

        # 并发配置
        self.max_workers = 15  # 同时翻译的线程数
        self.max_retries = 5  # 最大重试次数
        self.retry_delay = 1  # 重试延迟（秒）

        # 创建UI
        self.create_widgets()
        
    def create_widgets(self):
        """创建界面组件"""
        # 主框架
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 配置网格权重
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(3, weight=1)
        
        # 文件选择区域
        file_frame = ttk.LabelFrame(main_frame, text="文件选择", padding="10")
        file_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        file_frame.columnconfigure(1, weight=1)
        
        ttk.Button(file_frame, text="选择Word文档", command=self.select_file).grid(
            row=0, column=0, padx=(0, 10)
        )
        
        self.file_label = ttk.Label(file_frame, text="未选择文件", foreground="gray")
        self.file_label.grid(row=0, column=1, sticky=(tk.W, tk.E))
        
        # 控制按钮区域
        control_frame = ttk.Frame(main_frame)
        control_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        
        self.translate_btn = ttk.Button(
            control_frame, text="开始翻译", command=self.start_translation, state=tk.DISABLED
        )
        self.translate_btn.grid(row=0, column=0, padx=(0, 10))
        
        self.export_btn = ttk.Button(
            control_frame, text="导出文档", command=self.export_document, state=tk.DISABLED
        )
        self.export_btn.grid(row=0, column=1)
        
        # 进度区域
        progress_frame = ttk.LabelFrame(main_frame, text="翻译进度", padding="10")
        progress_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        progress_frame.columnconfigure(0, weight=1)
        
        self.progress_bar = ttk.Progressbar(
            progress_frame, mode='determinate', length=300
        )
        self.progress_bar.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 5))
        
        self.progress_label = ttk.Label(progress_frame, text="等待开始...")
        self.progress_label.grid(row=1, column=0, sticky=tk.W)
        
        # 翻译内容显示区域
        content_frame = ttk.LabelFrame(main_frame, text="翻译内容预览", padding="10")
        content_frame.grid(row=3, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        content_frame.columnconfigure(0, weight=1)
        content_frame.rowconfigure(0, weight=1)
        
        self.content_text = scrolledtext.ScrolledText(
            content_frame, wrap=tk.WORD, height=20, font=("Arial", 10)
        )
        self.content_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

    def translate_with_retry(self, text, src='en', dest='zh-cn'):
        """带重试机制的翻译方法"""
        for attempt in range(self.max_retries):
            try:
                # 为每次尝试创建新的翻译器实例，避免连接问题
                translator = Translator()
                translation = translator.translate(text, src=src, dest=dest)
                return translation.text
            except Exception as e:
                if attempt < self.max_retries - 1:
                    # 指数退避：每次重试延迟时间加倍
                    delay = self.retry_delay * (2 ** attempt)
                    self.show_current_text(f"[重试 {attempt + 1}/{self.max_retries}] 翻译失败，{delay}秒后重试...\n")
                    time.sleep(delay)
                else:
                    # 最后一次尝试失败，抛出异常
                    raise Exception(f"翻译失败（已重试{self.max_retries}次）: {str(e)}")

    def translate_paragraph(self, index, para_text):
        """翻译单个段落（用于并发处理）"""
        try:
            translated_text = self.translate_with_retry(para_text)
            return {
                'index': index,
                'original': para_text,
                'translated': translated_text,
                'success': True,
                'error': None
            }
        except Exception as e:
            return {
                'index': index,
                'original': para_text,
                'translated': None,
                'success': False,
                'error': str(e)
            }

    def select_file(self):
        """选择Word文档"""
        filename = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        
        if filename:
            self.input_file = filename
            self.file_label.config(text=os.path.basename(filename), foreground="black")
            self.translate_btn.config(state=tk.NORMAL)
            self.content_text.delete(1.0, tk.END)
            self.content_text.insert(tk.END, f"已选择文件: {filename}\n\n点击'开始翻译'按钮开始处理...\n")
            
    def start_translation(self):
        """开始翻译（在新线程中执行）"""
        if self.is_translating:
            messagebox.showwarning("警告", "翻译正在进行中，请稍候...")
            return
            
        if not self.input_file:
            messagebox.showerror("错误", "请先选择Word文档！")
            return
        
        # 禁用按钮
        self.translate_btn.config(state=tk.DISABLED)
        self.export_btn.config(state=tk.DISABLED)
        self.is_translating = True
        
        # 清空显示区域
        self.content_text.delete(1.0, tk.END)
        
        # 在新线程中执行翻译
        thread = threading.Thread(target=self.translate_document)
        thread.daemon = True
        thread.start()
        
    def translate_document(self):
        """翻译文档的主要逻辑（使用并发处理）"""
        try:
            # 读取Word文档
            self.update_ui("正在读取文档...", 0)
            doc = Document(self.input_file)

            # 保存原始文档引用，用于保留格式
            self.original_doc = doc

            # 获取所有需要翻译的段落（包含索引和原始段落对象）
            paragraph_data = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:  # 只翻译非空段落
                    paragraph_data.append({
                        'doc_index': i,  # 在文档中的位置
                        'para_obj': para,  # 原始段落对象
                        'text': text
                    })

            total = len(paragraph_data)

            if total == 0:
                self.update_ui("文档中没有可翻译的内容！", 0)
                self.is_translating = False
                self.root.after(0, lambda: self.translate_btn.config(state=tk.NORMAL))
                return

            self.update_ui(f"共找到 {total} 个段落，使用 {self.max_workers} 个线程并发翻译...", 0)
            self.show_current_text(f"\n开始并发翻译，线程数: {self.max_workers}，最大重试次数: {self.max_retries}\n\n")

            # 使用线程池并发翻译
            completed_count = 0
            results = {}

            with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                # 提交所有翻译任务
                future_to_index = {
                    executor.submit(self.translate_paragraph, i, item['text']): i
                    for i, item in enumerate(paragraph_data)
                }

                # 处理完成的任务
                for future in as_completed(future_to_index):
                    result = future.result()
                    results[result['index']] = result
                    completed_count += 1

                    # 更新进度
                    progress = int(completed_count / total * 100)
                    status = f"已完成 {completed_count}/{total} 段..."
                    self.update_ui(status, progress)

                    # 显示翻译结果
                    if result['success']:
                        self.show_current_text(
                            f"[完成 {result['index'] + 1}] ✓\n"
                            f"原文: {result['original'][:50]}...\n"
                            f"译文: {result['translated'][:50]}...\n\n"
                        )
                    else:
                        self.show_current_text(
                            f"[失败 {result['index'] + 1}] ✗\n"
                            f"原文: {result['original'][:50]}...\n"
                            f"错误: {result['error']}\n\n"
                        )

            # 在原文档基础上直接插入翻译内容，保留所有元素（图表、表格等）
            success_count = 0
            failed_count = 0

            # 创建翻译结果映射（从文档索引到翻译结果）
            translation_map = {}
            for i, item in enumerate(paragraph_data):
                translation_map[item['doc_index']] = results[i]

            # 直接使用原文档，在每个段落后插入翻译
            # 从后往前处理，避免插入新段落后索引变化
            doc_paragraphs = list(self.original_doc.paragraphs)

            # 需要从后往前处理，因为插入会改变后续段落的索引
            for doc_idx in range(len(doc_paragraphs) - 1, -1, -1):
                if doc_idx in translation_map:
                    result = translation_map[doc_idx]
                    para = doc_paragraphs[doc_idx]

                    # 在原段落之后插入翻译段落
                    if result['success']:
                        # 找到当前段落在文档中的位置
                        para_element = para._element
                        parent = para_element.getparent()

                        # 创建新的段落元素
                        new_para_element = OxmlElement('w:p')

                        # 创建段落属性
                        pPr = OxmlElement('w:pPr')
                        new_para_element.append(pPr)

                        # 创建run元素
                        new_run_element = OxmlElement('w:r')

                        # 创建run属性
                        rPr = OxmlElement('w:rPr')

                        # 设置中文字体
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'SimSun')
                        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'SimSun')
                        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'SimSun')
                        rPr.append(rFonts)

                        new_run_element.append(rPr)

                        # 创建文本元素
                        text_element = OxmlElement('w:t')
                        text_element.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                        text_element.text = result['translated']

                        new_run_element.append(text_element)
                        new_para_element.append(new_run_element)

                        # 在原段落后插入新段落
                        parent.insert(parent.index(para_element) + 1, new_para_element)

                        success_count += 1
                    else:
                        # 翻译失败，插入错误信息
                        para_element = para._element
                        parent = para_element.getparent()

                        new_para_element = OxmlElement('w:p')

                        pPr = OxmlElement('w:pPr')
                        new_para_element.append(pPr)

                        new_run_element = OxmlElement('w:r')
                        rPr = OxmlElement('w:rPr')

                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'SimSun')
                        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'SimSun')
                        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'SimSun')
                        rPr.append(rFonts)

                        new_run_element.append(rPr)

                        text_element = OxmlElement('w:t')
                        text_element.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                        text_element.text = f"[翻译失败: {result['error']}]"

                        new_run_element.append(text_element)
                        new_para_element.append(new_run_element)

                        parent.insert(parent.index(para_element) + 1, new_para_element)

                        failed_count += 1

            # 将修改后的原文档保存为结果
            self.translated_doc = self.original_doc

            # 翻译完成
            summary = f"翻译完成！成功: {success_count}，失败: {failed_count}。可以导出文档了。"
            self.update_ui(summary, 100)
            self.show_current_text(f"\n=== {summary} ===\n\n")

            # 启用导出按钮
            self.root.after(0, lambda: self.export_btn.config(state=tk.NORMAL))

        except Exception as e:
            error_msg = f"处理文档时出错: {str(e)}"
            self.update_ui(error_msg, 0)
            messagebox.showerror("错误", error_msg)

        finally:
            self.is_translating = False
            self.root.after(0, lambda: self.translate_btn.config(state=tk.NORMAL))
    
    def update_ui(self, message, progress):
        """更新UI显示（线程安全）"""
        self.root.after(0, lambda: self.progress_label.config(text=message))
        self.root.after(0, lambda: self.progress_bar.config(value=progress))
        
    def show_current_text(self, text):
        """显示当前翻译的文本（线程安全）"""
        def update():
            self.content_text.insert(tk.END, text)
            self.content_text.see(tk.END)
        self.root.after(0, update)
        
    def export_document(self):
        """导出翻译后的文档"""
        if not self.translated_doc:
            messagebox.showwarning("警告", "没有可导出的文档！")
            return
        
        # 生成默认文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_name = f"translated_{timestamp}.docx"
        
        # 选择保存位置
        filename = filedialog.asksaveasfilename(
            title="保存翻译文档",
            defaultextension=".docx",
            initialfile=default_name,
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        
        if filename:
            try:
                self.translated_doc.save(filename)
                messagebox.showinfo("成功", f"文档已成功保存到:\n{filename}")
                self.content_text.insert(tk.END, f"\n文档已导出: {filename}\n")
            except Exception as e:
                messagebox.showerror("错误", f"保存文档时出错:\n{str(e)}")


def main():
    root = tk.Tk()
    app = WordTranslatorApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()